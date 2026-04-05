# Import lightweight/standard libraries
from AppOpener import close, open as appopen
from webbrowser import open as webopen
from pywhatkit import search, playonyt
from dotenv import dotenv_values
from rich import print
import webbrowser
import subprocess
import requests
import asyncio
import shutil
import os
import datetime
import json

# --- SKILL IMPORTS ---
from Backend.FileOrganizer import CleanUp
from Backend.ImageGeneration import GenerateImages
from Backend.ClipboardReader import ReadClipboard
from Backend.Memory import Remember, Recall, Forget
from Backend.Security import SentryMode
from Backend.PhoneLink import SendToPhone  
import Backend.NervousSystem as Brain

# --- NEW SPOTIFY API CONTROLLER ---
from Backend.SpotifyController import PlaySpotifyTrack

# Load environment variables
env_vars = dotenv_values(".env")
GroqAPIKey = env_vars.get("GroqAPIKey")

messages = []
SystemChatBot = [{"role": "system", "content": f"Hello, I am {os.environ.get('Assistantname', 'Jarvis')}."}]

# Ensure Data & Workspace folders exist
os.makedirs("Data", exist_ok=True)
WORKSPACE_DIR = os.path.join(os.getcwd(), "Workspace")
os.makedirs(WORKSPACE_DIR, exist_ok=True)

# ==========================================================================================
# 1. PUBLIC API SKILLS (The "Free Forever" Pack)
# ==========================================================================================
def GetDefinition(word):
    try:
        url = f"https://api.dictionaryapi.dev/api/v2/entries/en/{word}"
        response = requests.get(url).json()
        definition = response[0]['meanings'][0]['definitions'][0]['definition']
        return f"The definition of {word} is: {definition}"
    except: return f"I couldn't find a definition for {word}."

def GetRandomFact():
    try: return requests.get("https://uselessfacts.jsph.pl/api/v2/facts/random").json()['text']
    except: return "I couldn't fetch a fact at the moment."

def GetSpaceImage():
    try:
        response = requests.get("https://api.nasa.gov/planetary/apod?api_key=DEMO_KEY").json()
        webbrowser.open(response['url']) 
        return f"Opening NASA's picture of the day: {response['title']}"
    except: return "NASA's satellites are not responding."

def GetActivity():
    try: return f"You could try to: {requests.get('https://bored-api.appbrewery.com/random').json()['activity']}."
    except: return "I suggest you take a nap."

def GetIPInfo():
    try:
        res = requests.get("http://ip-api.com/json").json()
        return f"You are currently located in {res['city']}, {res['country']}."
    except: return "I cannot trace your location, sir."

def GuessIdentity(name):
    try:
        age = requests.get(f"https://api.agify.io/?name={name}").json()['age']
        gender = requests.get(f"https://api.genderize.io/?name={name}").json()['gender']
        return f"I predict that {name} is a {age} year old {gender}."
    except: return "I couldn't analyze that name."

def GetAdvice():
    try: return requests.get("https://api.adviceslip.com/advice").json()['slip']['advice']
    except: return "Always trust your instincts."

def GetTechJargon():
    try: return requests.get("https://techy-api.vercel.app/api/json").json()['message']
    except: return "Have you tried turning it off and on again?"

def GetRandomJoke():
    try:
        res = requests.get("https://official-joke-api.appspot.com/random_joke").json()
        return f"{res['setup']} ... {res['punchline']}"
    except: return "I couldn't fetch a joke."

def GetWeatherPublic(city="Karachi"):
    try:
        temp = requests.get("https://api.open-meteo.com/v1/forecast?latitude=24.86&longitude=67.00&current_weather=true").json()['current_weather']['temperature']
        return f"The current temperature in {city} is {temp} degrees Celsius."
    except: return "Weather data unavailable."

def GetCryptoPrice(coin):
    try:
        price = requests.get(f"https://api.coingecko.com/api/v3/simple/price?ids={coin}&vs_currencies=usd").json()[coin]['usd']
        return f"The price of {coin} is ${price}."
    except: return f"Check the ticker symbol again."

# ==========================================================================================
# 2. SYSTEM CONTROLS
# ==========================================================================================
def VolumeControl(command):
    import pyautogui
    if "up" in command or "increase" in command: pyautogui.press("volumeup", presses=5)
    elif "down" in command or "decrease" in command: pyautogui.press("volumedown", presses=5)
    elif "mute" in command or "silent" in command: pyautogui.press("volumemute")
    return "Volume Adjusted"

def BrightnessControl(command):
    try:
        import screen_brightness_control as sbc
        current = sbc.get_brightness()
        if not current: return "Brightness control not supported."
        current_val = current[0]
        if "up" in command: sbc.set_brightness(min(current_val + 20, 100))
        elif "down" in command: sbc.set_brightness(max(current_val - 20, 0))
        elif "set" in command:
            import re
            numbers = re.findall(r'\d+', command)
            if numbers: sbc.set_brightness(int(numbers[0]))
        return "Brightness Adjusted"
    except: return "Error adjusting brightness."

def SystemStats():
    import psutil
    battery = psutil.sensors_battery()
    if battery:
        status = "Plugged In" if battery.power_plugged else "Running on Battery"
        return f"Battery is at {battery.percent}% and {status}."
    return "System is running on AC Power."

def TakeScreenshot():
    import pyautogui
    path = os.path.join("Data", f"Screenshot_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.png")
    pyautogui.screenshot().save(path)
    os.startfile(path)
    return "Screenshot saved."

def TakePhoto():
    import cv2
    cap = cv2.VideoCapture(0)
    if not cap.isOpened(): return "No Camera."
    ret, frame = cap.read()
    if ret:
        path = os.path.join("Data", f"Camera_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg")
        cv2.imwrite(path, frame)
        cap.release()
        os.startfile(path)
        return "Photo taken."
    cap.release()
    return "Camera Error."

def PowerControl(command):
    if "shutdown" in command: os.system("shutdown /s /t 5")
    elif "restart" in command: os.system("shutdown /r /t 5")
    elif "sleep" in command: os.system("rundll32.exe powrprof.dll,SetSuspendState 0,1,0")
    return "Power command initiated."

# ==========================================================================================
# 3. CORE AUTOMATION & MS WORD GENERATION
# ==========================================================================================
def GoogleSearch(Topic): 
    search(Topic)
    return f"Searching Google for {Topic}, sir."

def PlayYoutube(query): 
    playonyt(query)
    return f"Playing {query} on YouTube, sir."

def YouTubeSearch(Topic): 
    webbrowser.open(f"https://www.youtube.com/results?search_query={Topic}")
    return f"Searching YouTube for {Topic}, sir."

def Content(Topic):
    """Generates professional content and outputs a native MS Word document."""
    from groq import Groq
    try:
        from docx import Document
    except ImportError:
        return "Sir, the python-docx library is missing. Please run 'pip install python-docx' in your terminal."

    client = Groq(api_key=GroqAPIKey)
    Topic = Topic.replace("content ", "").replace("write ", "").strip()

    try:
        with open(os.path.join("Data", "ChatLog.json"), "r", encoding="utf-8") as f:
            chat_history = json.load(f)[-10:] 
    except:
        chat_history = []

    chat_history.append({"role": "user", "content": f"Please generate a detailed, professionally formatted document about: {Topic}. Do not include markdown asterisks, just clean text with clear paragraphs."})
    
    completion = client.chat.completions.create(
        model="llama-3.3-70b-versatile", 
        messages=SystemChatBot + chat_history, 
        max_tokens=2048
    )
    Answer = completion.choices[0].message.content
    
    # Generate the MS Word Document
    doc = Document()
    doc.add_heading(Topic.title(), 0)
    
    for paragraph in Answer.split('\n\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    
    safe_title = "".join(c for c in Topic[:20] if c.isalnum() or c in (' ', '_')).replace(' ', '_')
    if not safe_title: safe_title = "Generated_Content"
    
    # Save directly to the Workspace
    file_path = os.path.join(WORKSPACE_DIR, f"{safe_title}.docx")
    doc.save(file_path)
    
    # Open natively in MS Word
    os.startfile(file_path)
    
    return f"I have written the report on {Topic} and opened it in Microsoft Word, sir."

# --- DYNAMIC FILE HANDLING ---
def CreateFolder(folder_name):
    target = os.path.join(WORKSPACE_DIR, folder_name.strip())
    os.makedirs(target, exist_ok=True)
    return f"Folder '{folder_name}' has been created in your Workspace, sir."

def DeleteWorkspaceItem(item_name):
    target = os.path.join(WORKSPACE_DIR, item_name.strip())
    if not os.path.exists(target):
        return f"Sir, I could not find '{item_name}' in the Workspace."
    try:
        if os.path.isdir(target):
            shutil.rmtree(target)
            return f"Folder '{item_name}' has been deleted."
        else:
            os.remove(target)
            return f"File '{item_name}' has been deleted."
    except Exception as e:
        return f"Error deleting item: {e}"

def OpenApp(app, sess=requests.session()):
    app_name = app.strip().lower().replace(".", "").replace(",", "").replace("please", "").strip()
    links = {
        "facebook": "https://www.facebook.com", "instagram": "https://www.instagram.com",
        "twitter": "https://www.twitter.com", "tiktok": "https://www.tiktok.com",
        "spotify": "https://open.spotify.com", "youtube": "https://www.youtube.com",
        "chatgpt": "https://chatgpt.com", "gmail": "https://mail.google.com"
    }
    if app_name in links: 
        webbrowser.open(links[app_name])
        return f"Opening {app_name}, sir."
        
    if any(ext in app_name for ext in [".com", ".org", ".net", ".edu", ".pk"]) or "http" in app_name:
        target = f"https://{app_name}" if "http" not in app_name else app_name
        webbrowser.open(target)
        return f"Navigating to {app_name} on your browser, sir."

    try: 
        appopen(app_name, match_closest=True, output=False, throw_error=True)
        return f"Opening {app_name} from your system, sir."
    except: 
        search_url = f"https://www.google.com/search?q={app_name}"
        webbrowser.open(search_url)
        return "You don't have this app, sir. I am opening information about it on Chrome."

def CloseApp(app):
    app_name = app.strip().lower().replace(".", "").replace(",", "").replace("please", "").strip()
    try: 
        close(app_name, match_closest=True, output=False, throw_error=True)
        return f"Closing {app_name}, sir."
    except: 
        return f"I couldn't find an active window for {app_name}, sir."

# ==========================================================================================
# 4. EXECUTION ENGINE
# ==========================================================================================
async def TranslateAndExecute(commands: list[str]):
    funcs = []
    for command in commands:
        if command.startswith("open "):
            fun = asyncio.to_thread(OpenApp, command.removeprefix("open "))
            funcs.append(fun)
        elif command.startswith("close "):
            fun = asyncio.to_thread(CloseApp, command.removeprefix("close "))
            funcs.append(fun)
        elif command.startswith("play "):
            if "spotify" in command.lower():
                # 🔥 NOW WIRED DIRECTLY TO THE SPOTIPY API MODULE
                fun = asyncio.to_thread(PlaySpotifyTrack, command.removeprefix("play "))
            else:
                fun = asyncio.to_thread(PlayYoutube, command.removeprefix("play "))
            funcs.append(fun)
        elif command.startswith("content ") or command.startswith("write "):
            fun = asyncio.to_thread(Content, command)
            funcs.append(fun)
        elif command.startswith("google search "):
            fun = asyncio.to_thread(GoogleSearch, command.removeprefix("google search "))
            funcs.append(fun)
        elif command.startswith("youtube search "):
            fun = asyncio.to_thread(YouTubeSearch, command.removeprefix("youtube search "))
            funcs.append(fun)
            
        elif command.startswith("create folder "):
            fun = asyncio.to_thread(CreateFolder, command.removeprefix("create folder "))
            funcs.append(fun)
        elif command.startswith("delete "):
            fun = asyncio.to_thread(DeleteWorkspaceItem, command.removeprefix("delete "))
            funcs.append(fun)
            
        elif command.startswith("system "):
            cmd = command.removeprefix("system ").lower()
            if "volume" in cmd or "mute" in cmd: yield VolumeControl(cmd)
            elif "brightness" in cmd: yield BrightnessControl(cmd)
            elif "screenshot" in cmd: yield TakeScreenshot()
            elif "photo" in cmd or "camera" in cmd: yield TakePhoto()
            elif "battery" in cmd or "stats" in cmd: yield SystemStats()
            elif "shutdown" in cmd or "restart" in cmd: yield PowerControl(cmd)
        else:
            print(f"No Function Found. For {command}")

    results = await asyncio.gather(*funcs)
    for result in results:
        yield result

async def Automation(commands: list[str]):
    for command in commands:
        query = str(command).lower()

        # --- NERVOUS SYSTEM TRIGGERS ---
        if "fact about" in query and any(char.isdigit() for char in query):
            num = ''.join(filter(str.isdigit, query))
            yield Brain.GetNumberFact(num); continue

        if "book" in query and "search" in query:
            q = query.replace("book", "").replace("search", "").strip()
            yield Brain.GetBookInfo(q); continue

        if "urban" in query or "slang" in query:
            q = query.replace("urban", "").replace("slang", "").replace("define", "").strip()
            yield Brain.GetUrbanDef(q); continue

        if "tech" in query and "news" in query:
            yield Brain.GetTechNews(); continue

        if "space" in query and "news" in query:
            yield Brain.GetSpaceNews(); continue

        if "iss" in query and "location" in query:
            yield Brain.GetISSLocation(); continue

        if "coordinates" in query or "where is" in query:
            city = query.replace("coordinates", "").replace("where is", "").strip()
            yield Brain.GetCoordinates(city); continue

        if "currency" in query or "exchange" in query:
            curr = query.split()[-1]
            yield Brain.GetCurrencyRate(curr); continue

        if "show" in query or "series" in query:
            q = query.replace("show", "").replace("series", "").replace("info", "").strip()
            yield Brain.GetMovieInfo(q); continue

        if "anime" in query:
            q = query.replace("anime", "").replace("info", "").strip()
            yield Brain.GetAnimeInfo(q); continue

        if "github" in query or "repo" in query:
            q = query.replace("github", "").replace("repo", "").strip()
            yield Brain.GetGithubTrending(q); continue
        
        # --- ORIGINAL SKILLS ---
        if "define" in query:
            word = query.replace("define", "").strip()
            yield GetDefinition(word); continue
        if "fact" in query: yield GetRandomFact(); continue
        if "joke" in query: yield GetRandomJoke(); continue
        if "advice" in query: yield GetAdvice(); continue
        if "bored" in query or "activity" in query: yield GetActivity(); continue
        if "tech" in query and "say" in query: yield GetTechJargon(); continue
        if "space" in query and "image" in query: yield GetSpaceImage(); continue
        if "where am i" in query or "location" in query: yield GetIPInfo(); continue
        if "guess" in query and ("age" in query or "gender" in query):
            name = query.split()[-1] 
            yield GuessIdentity(name); continue
        
        if "weather" in query: yield GetWeatherPublic(); continue
        if "price" in query:
            coin = "bitcoin"
            if "ethereum" in query: coin = "ethereum"
            elif "doge" in query: coin = "dogecoin"
            elif "solana" in query: coin = "solana"
            yield GetCryptoPrice(coin); continue
            
        # --- ADVANCED SKILLS ---
        if "cleanup" in query or ("organize" in query and "files" in query):
            yield CleanUp(); continue
            
        if "generate" in query and "image" in query:
            prompt = query.replace("generate", "").replace("image", "").strip()
            yield GenerateImages(prompt); continue
            
        if "clipboard" in query or ("read" in query and "copy" in query):
            yield ReadClipboard(query); continue

        if "remember" in query: yield Remember(query); continue
        if "recall" in query or "what do you remember" in query: yield Recall(); continue
        if "forget" in query: yield Forget(query); continue

        if "sentry" in query or "security" in query or "jarvis" in query:
            yield SentryMode(); continue
            
        if "send to phone" in query or "phone link" in query:
            yield SendToPhone(query); continue

    # --- REGULAR COMMANDS (Apps, Web, System, Files, Word) ---
    async for result in TranslateAndExecute(commands):
        yield result